from docx import Document

class DocumentProcessor:
    def __init__(self, file_path):
        self.doc = Document(file_path)

    def get_text(self):
        full_text = []
        for paragraph in self.doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)

    def get_snippets(self, query):
        snippets = []
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            if query.lower() in text.lower():
                snippets.append(text)
        return snippets
